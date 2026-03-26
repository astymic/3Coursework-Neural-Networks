"""
Скрипт генерації .docx звітів для ЛР 1-5.
Оформлення: Times New Roman 14, полуторний інтервал, абзацний відступ 1.25 см,
вирівнювання по ширині, заголовки по центру жирним 16 пт.
"""
import os
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
TITLE_PATH = os.path.join(BASE, 'Title.docx')

def insert_title_page(doc, lab_number):
    """Вставляє титульну сторінку з Title.docx на початок документа, змінюючи номер ЛР."""
    title_doc = Document(TITLE_PATH)

    # Копіюємо секційні налаштування (поля сторінки) з Title.docx
    # для першої секції, якщо потрібно

    # Збираємо всі елементи body з шаблону титулки
    title_elements = []
    for element in title_doc.element.body:
        title_elements.append(copy.deepcopy(element))

    # Додаємо розрив сторінки після титулки
    page_break_p = OxmlElement('w:p')
    page_break_r = OxmlElement('w:r')
    page_break_br = OxmlElement('w:br')
    page_break_br.set(qn('w:type'), 'page')
    page_break_r.append(page_break_br)
    page_break_p.append(page_break_r)
    title_elements.append(page_break_p)

    # Вставляємо елементи на початок документа (перед усім існуючим контентом)
    body = doc.element.body
    first_child = body[0] if len(body) > 0 else None

    for elem in title_elements:
        if first_child is not None:
            body.insert(body.index(first_child), elem)
        else:
            body.append(elem)

    # Замінюємо номер ЛР у титулці
    for paragraph in doc.paragraphs:
        if 'лабораторної роботи' in paragraph.text.lower():
            for run in paragraph.runs:
                if '№' in run.text:
                    # Замінюємо будь-який номер після № на потрібний
                    import re
                    run.text = re.sub(r'№\d+', f'№{lab_number}', run.text)
            break  # Лише перше входження (титулка)

# ───── Утиліти форматування ─────
def set_cell_shading(cell, color_hex):
    """Зафарбовує комірку таблиці."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def setup_styles(doc):
    """Налаштовує базовий стиль документа."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Поля сторінки
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

def add_heading_centered(doc, text, level=1, font_size=16):
    """Додає заголовок по центру, жирний."""
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.first_line_indent = Cm(0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text.upper() if level == 1 else text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)

def add_body(doc, text):
    """Додає абзац основного тексту."""
    p = doc.add_paragraph(text)
    return p

def add_bold_item(doc, label, value):
    """Додає абзац з жирним полем та звичайним значенням."""
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r2 = p.add_run(value)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    return p

def add_code_block(doc, code_text, lang_label='Python'):
    """Додає блок коду (Courier New 10pt, рамка, сіре тло)."""
    # Підпис
    lbl = doc.add_paragraph()
    lbl.paragraph_format.first_line_indent = Cm(0)
    lbl.paragraph_format.space_before = Pt(6)
    r = lbl.add_run(f'Лістинг ({lang_label}):')
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    # Код
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

def add_image_with_caption(doc, image_path, caption, width_cm=15):
    """Додає зображення та підпис під ним."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.space_after = Pt(6)
        r = cap.add_run(caption)
        r.italic = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    else:
        add_body(doc, f'[Файл зображення не знайдено: {image_path}]')

def read_code_file(path):
    """Зчитує текстовий файл."""
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


# ─────────────────────────────────────────────
# ЛР 1
# ─────────────────────────────────────────────
def generate_lab1():
    doc = Document()
    setup_styles(doc)
    lr_dir = os.path.join(BASE, 'ЛР1')

    add_heading_centered(doc, 'Звіт: Лабораторна робота №1')
    add_heading_centered(doc, 'Одношаровий Персептрон', level=2, font_size=14)

    # --- Мета ---
    add_heading_centered(doc, '1. Архітектура мережі та гіперпараметри ШНМ', level=2, font_size=14)
    add_bold_item(doc, 'Тип мережі: ', 'Одношаровий перцептрон (1 нейрон).')
    add_bold_item(doc, 'Кількість вхідних вузлів: ', '3 (Дохід, Кредитний рейтинг, Відношення боргу до доходу).')
    add_bold_item(doc, 'Функція активації: ', 'Сигмоїдна σ(x) = 1 / (1 + e^(-x)).')
    add_bold_item(doc, 'Швидкість навчання (Learning Rate): ', '0.05.')
    add_bold_item(doc, 'Початкові вагові коефіцієнти: ', 'Згенеровані випадковим чином в інтервалі [0, 1].')
    add_bold_item(doc, 'Кількість епох навчання: ', '10, 50, 100, 1000.')

    # --- Алгоритм вибірки ---
    add_heading_centered(doc, '2. Алгоритм побудови навчальної та тестової вибірки', level=2, font_size=14)
    add_body(doc, 'Вибірки були згенеровані через Python (numpy.random.rand(num_samples)), що одразу забезпечило їх нормалізацію в межах [0, 1]. Ми імітували параметри для кредитного скорингу банку:')
    add_body(doc, '• Income_Norm (Дохід)')
    add_body(doc, '• CreditScore_Norm (Кредитний рейтинг)')
    add_body(doc, '• DTI_Ratio_Norm (Борг до доходу – Debt-To-Income)')
    add_body(doc, 'Вихідний результат (видати чи відхилити кредит) формувався за власною нелінійною функцією: output = 1, якщо (Income_Norm + CreditScore_Norm - DTI_Ratio_Norm) > 0.6 ; в іншому випадку output = 0.')
    add_body(doc, 'Було сформовано 500 прикладів для train_data.xlsx та 100 прикладів для test_data.xlsx.')

    # --- Алгоритм навчання ---
    add_heading_centered(doc, '3. Алгоритм навчання ШНМ', level=2, font_size=14)
    add_body(doc, 'Навчання мережі здійснюється за дельта-правилом із використанням швидкості навчання та похідної сигмоїди.')
    add_body(doc, 'Для кожної епохи, для кожного прикладу з навчальної вибірки:')
    add_body(doc, '1. Зчитується сума добутків входів на їх ваги: sum = Σ (X[i] * W[i])')
    add_body(doc, '2. Обчислюється результат активації: output = sigmoid(sum)')
    add_body(doc, '3. Обчислюється помилка: error = Expected_Output - output')
    add_body(doc, '4. Вектор ваг оновлюється за правилом: ΔW[j] = learning_rate * error * X[i] * (output * (1 - output))')

    # --- Код ---
    add_heading_centered(doc, '4. Текст програми', level=2, font_size=14)
    code = read_code_file(os.path.join(lr_dir, 'Lab1_Perceptron.py'))
    add_code_block(doc, code)

    # --- Результати ---
    add_heading_centered(doc, '5. Результати навчання', level=2, font_size=14)
    results = [
        ('10 епох', '0.3228', '87.00%', '[1.312, 1.635, -3.052]'),
        ('50 епох', '0.2295', '86.00%', '[2.820, 2.992, -5.921]'),
        ('100 епох', '0.2075', '86.00%', '[3.418, 3.594, -7.115]'),
        ('1000 епох', '0.1808', '86.00%', '[4.506, 4.717, -9.297]'),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Кількість епох', 'Мін. помилка', 'Точність', 'Ваги']):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
    for row_data in results:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)

    add_image_with_caption(doc, os.path.join(lr_dir, 'training_plot.png'), 'Рисунок 1 – Графік помилки під час навчання')

    # --- Висновки ---
    add_heading_centered(doc, '6. Висновки', level=2, font_size=14)
    add_body(doc, 'Перцептрон з 1 нейроном зміг вивчити власну лінійну залежність навчального датасету і досяг точності 86-87% на перевірочному датасеті (тестовій вибірці).')
    add_body(doc, 'Сигнатура вагових коефіцієнтів повністю відповідає логіці нашої кастомної функції генерації output: ваги для Income та CreditScore є позитивними, а вага для DTI_Ratio — різко негативною. Нейромережа сама зрозуміла, що високий відсоток боргу впливає негативно на видачу кредиту.')
    add_body(doc, 'Вплив кількості епох: Після перших 20-50 епох середня помилка різко падала. Далі мережа "довивчала" паттерни. Однак точність на тестовому сеті (86%) завмерла, незважаючи на падіння похибки на навчальному сеті (через неможливість 100% лінійної сепарації 3-мірного простору).')
    add_body(doc, 'Сигмоїдна функція дозволила зробити м\'яке навчання ваг порівняно з жорсткою ступінчастою функцією, і показала чудову стабільність градієнту.')

    insert_title_page(doc, 1)
    out = os.path.join(lr_dir, 'Lab1_Report.docx')
    doc.save(out)
    print(f'[OK] {out}')


# ─────────────────────────────────────────────
# ЛР 2
# ─────────────────────────────────────────────
def generate_lab2():
    doc = Document()
    setup_styles(doc)
    lr_dir = os.path.join(BASE, 'ЛР2')

    add_heading_centered(doc, 'Звіт: Лабораторна робота №2')
    add_heading_centered(doc, 'Багатошарові Мережі (MNIST / Fashion MNIST)', level=2, font_size=14)

    add_heading_centered(doc, 'Мета роботи', level=2, font_size=14)
    add_body(doc, 'Написати програму мовою Python, що створює та навчає нейронну мережу для розпізнавання рукописних цифр (MNIST) та графічних символів (Fashion MNIST) за допомогою NumPy.')

    # --- Код ---
    add_heading_centered(doc, '1. Коди програми Network2 із коментарями', level=2, font_size=14)
    add_body(doc, 'Нижче наведено повний лістинг коду для розпізнавання Fashion MNIST зі зміненою функцією активації (ReLU) та ініціалізацією ваг за методом He, що дозволило уникнути проблеми згасаючого/вибухаючого градієнта.')
    code = read_code_file(os.path.join(lr_dir, 'Lab2_Network2.py'))
    add_code_block(doc, code)

    # --- Скріншоти ---
    add_heading_centered(doc, '2. Результати навчання та класифікації', level=2, font_size=14)
    add_body(doc, 'Швидкість навчання (eta) була встановлена на 0.5, кількість епох — 10, розмір міні-батчу — 10. Точність на тестовій вибірці стабільно досягла ~85.5% (8539 правильних прогнозів з 10000).')
    add_image_with_caption(doc, os.path.join(lr_dir, 'Network2_Visual_Test.png'), 'Рисунок 1 – Результат класифікації випадкових зображень Fashion MNIST')

    # --- Висновки ---
    add_heading_centered(doc, '3. Висновки', level=2, font_size=14)
    add_body(doc, '1. Базова мережа (Sigmoid + MNIST): Базовий багатошаровий перцептрон з активаційною функцією Sigmoid зміг досягти високої точності ~94.4% на датасеті MNIST (рукописні цифри) після 10 епох. Дані завантажувались локально з файлу mnist.pkl.gz за допомогою бібліотеки pickle, що підтверджує можливість не залежати від зовнішніх API keras.')
    add_body(doc, '2. Модифікована мережа (ReLU + Fashion MNIST): Заміна сигмоїди на ReLU в прихованому шарі потребувала зміни похідної на relu_prime у алгоритмі зворотного поширення помилки (backprop).')
    add_body(doc, '3. Особливості ReLU: ReLU дозволив прискорити прямий і зворотній прохід завдяки відсутності складної експоненти. Однак, ReLU виявився набагато чутливішим до гіперпараметрів і ініціалізації ваг. З високим learning_rate = 3.0 нейрони ставали "мертвими" (повертали нуль), тому швидкість навчання довелося зменшити до орієнтовно 0.5, а для ініціалізації ваг використати метод He Initialization.')
    add_body(doc, '4. Результат розпізнавання: Модифікована мережа розпізнала 85.5% одягу з бібліотеки Fashion MNIST, і успішно провела класифікацію випадкових зображень під час процедури візуального тестування (як продемонстровано на графіках matplotlib).')

    insert_title_page(doc, 2)
    out = os.path.join(lr_dir, 'Lab2_Report.docx')
    doc.save(out)
    print(f'[OK] {out}')


# ─────────────────────────────────────────────
# ЛР 3
# ─────────────────────────────────────────────
def generate_lab3():
    doc = Document()
    setup_styles(doc)
    lr_dir = os.path.join(BASE, 'ЛР3')

    add_heading_centered(doc, 'Звіт: Лабораторна робота №3')
    add_heading_centered(doc, 'Згорткова нейронна мережа (CNN)', level=2, font_size=14)

    add_heading_centered(doc, 'Мета роботи', level=2, font_size=14)
    add_body(doc, 'Ознайомлення зі створенням та навчанням згорткової нейронної мережі (CNN) з використанням TensorFlow/Keras та візуалізація роботи за допомогою Matplotlib на базі датасету MNIST. З урахуванням індивідуальних завдань.')

    # --- Код ---
    add_heading_centered(doc, '1. Код програми з коментарями', level=2, font_size=14)
    add_body(doc, 'Повний лістинг з включенням індивідуальних завдань:')
    code = read_code_file(os.path.join(lr_dir, 'Lab3_CNN.py'))
    add_code_block(doc, code)
    add_body(doc, 'Після підготовки моделі було виконано навчання (на GPU) для 10, 50 та 100 епох. Оскільки функція втрат змінена на mean_squared_error, мітки класів (0-9) були попередньо перетворені у формат One-Hot Encoding за допомогою функції to_categorical().')

    # --- Графіки ---
    add_heading_centered(doc, '2. Графіки результатів тренування мережі', level=2, font_size=14)
    add_image_with_caption(doc, os.path.join(lr_dir, 'Lab3_CNN_Training_Plots.png'), 'Рисунок 1 – Графіки Accuracy та Loss (MSE) для 10, 50 та 100 епох')

    add_heading_centered(doc, 'Підсумкові показники на тестовій вибірці', level=2, font_size=14)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Кількість епох', 'Test Accuracy', 'Test MSE']):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
    for row_data in [('10', '0.8793', '0.0216'), ('50', '0.9316', '0.0108'), ('100', '0.9531', '0.0075')]:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)

    # --- Висновки ---
    add_heading_centered(doc, '3. Оцінка результатів та висновки', level=2, font_size=14)
    add_body(doc, '1. Зміна архітектури (MaxPooling2D): Додавання шару Макспулінгу значно зменшило розмірність карти ознак (від 26×26 до 13×13). Це різко скоротило кількість параметрів наступного шару Dense(64), прискорило навчання та зробило мережу більш стійкою до незначних зміщень цифр на зображенні.')
    add_body(doc, '2. Оптимізатор (SGD): Стохастичний градієнтний спуск (SGD) виявився повільнішим за сучасні алгоритми на зразок Adam. Замість швидкої збіжності за перші 3-5 епох, SGD стабільно та поступово зменшував похибку аж до сотої епохи, не досягаючи ефекту "завмирання" занадто рано.')
    add_body(doc, '3. Функція втрат (mean_squared_error): Стандартно в задачах багатокласової класифікації з softmax використовується categorical_crossentropy. Зміна на MSE вимагала обов\'язкового форматування міток у One-Hot вектори, щоб рахувати відстань між лінійно незалежними категоріями. MSE працює для цієї задачі, але її градієнти були меншими в порівнянні з ентропією, через що навчання було більш плавним.')
    add_body(doc, '4. Кількість епох (10, 50, 100): Чітко простежується закономірність: зі збільшенням кількості епох від 10 до 100, точність (Accuracy) лінійно зростала з 87% до 95.3%, а похибка (MSE) зменшилась утричі — з 0.0216 до 0.0075. При цьому перенавчання (Overfitting) на графіках не спостерігається — похибка на тестовій вибірці падала синхронно з похибкою на навчальній. Це ідеальна поведінка.')

    insert_title_page(doc, 3)
    out = os.path.join(lr_dir, 'Lab3_Report.docx')
    doc.save(out)
    print(f'[OK] {out}')


# ─────────────────────────────────────────────
# ЛР 4
# ─────────────────────────────────────────────
def generate_lab4():
    doc = Document()
    setup_styles(doc)
    lr_dir = os.path.join(BASE, 'ЛР4')

    add_heading_centered(doc, 'Звіт: Лабораторна робота №4')
    add_heading_centered(doc, 'Рекурентна Нейронна Мережа (RNN)', level=2, font_size=14)

    add_heading_centered(doc, 'Мета роботи', level=2, font_size=14)
    add_body(doc, 'Вивчити побудову та навчання Рекурентної Нейронної Мережі (RNN) мовою Python з використанням бібліотеки numpy "з нуля" (архітектура "багато до одного").')

    add_heading_centered(doc, '1. Формулювання обраного завдання класифікації', level=2, font_size=14)
    add_body(doc, 'Для індивідуального завдання було обрано задачу: Класифікація фраз технічного чи ліричного тексту.')
    add_body(doc, 'Мережа має отримувати на вхід послідовність слів (фразу з 3 слів) і визначати її приналежність до одного з двох класів:')
    add_body(doc, '• 0 – Технічний текст (алгоритми, код, сервери).')
    add_body(doc, '• 1 – Ліричний текст (любов, весна, небо).')

    add_heading_centered(doc, '2. Створені набори даних (Словники/Датасети)', level=2, font_size=14)
    add_body(doc, 'Оскільки зовнішні готові датасети не використовувалися (щоб максимізувати роботу "з нуля"), я самостійно згенерував словник:')
    add_bold_item(doc, 'Технічні слова: ', "'алгоритм', 'дані', 'функція', 'масив', 'сервер', 'база', 'компілятор', 'код', 'змінна', 'інтерфейс'.")
    add_bold_item(doc, 'Ліричні слова: ', "'любов', 'душа', 'сонце', 'вітер', 'серце', 'весна', 'сльоза', 'радість', 'небо', 'мрія'.")
    add_body(doc, 'Було згенеровано 1000 тренувальних фраз (по 3 випадкових слова з відповідного класу) та 200 тестових фраз. Всі слова конвертувалися у вектори One-Hot Encoding (довжина вектора = 20 слів).')

    add_heading_centered(doc, '3. Вибрана архітектура мережі та гіперпараметри', level=2, font_size=14)
    add_bold_item(doc, 'Архітектура: ', 'Vanilla RNN, багато до одного.')
    add_bold_item(doc, 'Кодування входу: ', 'One-Hot (20 розмірність).')
    add_bold_item(doc, 'Прихований шар: ', '64 нейрони.')
    add_bold_item(doc, 'Вихідний шар: ', '1 нейрон для бінарної класифікації.')
    add_bold_item(doc, 'Функція активації прихованого шару: ', 'tanh. Класичний вибір для RNN для розподілу виходів в діапазоні [-1, 1].')
    add_bold_item(doc, 'Функція кросу (Cost): ', 'Binary Cross Entropy loss.')
    add_bold_item(doc, 'Оновлення ваг: ', 'Backpropagation Through Time (BPTT).')
    add_bold_item(doc, 'Кількість епох: ', '100.')
    add_bold_item(doc, 'Швидкість навчання: ', '0.05. Було застосовано Gradient Clipping [-1, 1] для стабілізації BPTT.')

    # --- Код ---
    add_heading_centered(doc, '4. Коди програм із коментарями', level=2, font_size=14)
    code = read_code_file(os.path.join(lr_dir, 'Lab4_RNN.py'))
    add_code_block(doc, code)

    # --- Результати ---
    add_heading_centered(doc, '5. Скріншоти та результати класифікації', level=2, font_size=14)
    add_body(doc, 'Оскільки RNN успішно засвоїла залежності словників, вона правильно zgенерувала всі класи для тестових прикладів:')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Фраза', 'Передбачення', 'Очікувалося']):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
    for row_data in [
        ('функція алгоритм код', 'Технічний (1.00)', 'Технічний'),
        ('небо серце любов', 'Ліричний (1.00)', 'Ліричний'),
        ('сервер база дані', 'Технічний (1.00)', 'Технічний'),
        ('радість весна душа', 'Ліричний (1.00)', 'Ліричний'),
    ]:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)

    add_image_with_caption(doc, os.path.join(lr_dir, 'Lab4_RNN_Metrics.png'), 'Рисунок 1 – Метрики Loss та Accuracy під час навчання RNN')
    add_body(doc, 'Точність на тестовій вибірці (200 фраз): 100.00%.')

    # --- Висновки ---
    add_heading_centered(doc, '6. Висновки', level=2, font_size=14)
    add_body(doc, 'При реалізації RNN з нуля я зіткнувся з класичною проблемою вибухаючих градієнтів під час зворотного поширення помилки через час (BPTT). Це призводило до переповнення типу (NaN) у функції Sigmoid та Tanh. Цю проблему було успішно вирішено 2 кроками:')
    add_body(doc, '1. Діленням ваг при початковій ініціалізації (np.random.randn() / 1000).')
    add_body(doc, '2. Додаванням Gradient Clipping (np.clip(d, -1, 1, out=d)) перед кожним оновленням ваг.')
    add_body(doc, 'У результаті повністю створена з нуля рекурентна нейронна мережа на базі numpy змогла безпомилково класифікувати тип тексту. Навчання на CPU зайняло лічені секунди через невеликий розмір датасету та архітектури.')

    insert_title_page(doc, 4)
    out = os.path.join(lr_dir, 'Lab4_Report.docx')
    doc.save(out)
    print(f'[OK] {out}')


# ─────────────────────────────────────────────
# ЛР 5
# ─────────────────────────────────────────────
def generate_lab5():
    doc = Document()
    setup_styles(doc)
    lr_dir = os.path.join(BASE, 'ЛР5')

    add_heading_centered(doc, 'Звіт: Лабораторна робота №5')
    add_heading_centered(doc, 'Навчання без вчителя (Кластеризація)', level=2, font_size=14)

    add_heading_centered(doc, 'Мета роботи', level=2, font_size=14)
    add_body(doc, 'Вивчити методи навчання без вчителя (k-means, ієрархічна кластеризація, t-SNE, DBSCAN) за допомогою бібліотеки scikit-learn.')

    # --- Код ---
    add_heading_centered(doc, 'Код програми', level=2, font_size=14)
    code = read_code_file(os.path.join(lr_dir, 'Lab5_Clustering.py'))
    add_code_block(doc, code)

    # --- 1. K-Means ---
    add_heading_centered(doc, '1. Метод k-середніх (k-means) на різних парах ознак', level=2, font_size=14)
    add_body(doc, 'Для дослідження поведінки кластеризації методом k-means було обрано дві нові пари ознак з датасету Iris:')
    add_body(doc, '1. Petal length (довжина пелюстки) та Petal width (ширина пелюстки).')
    add_body(doc, '2. Sepal length (довжина чашолистка) та Petal length.')
    add_body(doc, 'Аналізуючи графік, можна побачити, що використання ознак пелюстки (Petal) дає значно чіткіший і щільніший розподіл на 3 кластери порівняно з ознаками чашолистка (Sepal). Ірис Setosa ідеально відокремлюється від інших двох видів у будь-якій комбінації, але Versicolor та Virginica мають невелике перекриття, з яким k-means справляється краще саме на ознаках пелюстки (за рахунок мінімізації дисперсії від центроїдів).')
    add_image_with_caption(doc, os.path.join(lr_dir, 'Lab5_KMeans_Clusters.png'), 'Рисунок 1 – K-means кластеризація на двох парах ознак')

    # --- 2. Ієрархічна ---
    add_heading_centered(doc, '2. Ієрархічна кластеризація та Дендрограма', level=2, font_size=14)
    add_body(doc, 'Було виконано агломеративну ієрархічну кластеризацію датасету Ірис Фішера з використанням методу Уорда (ward), який мінімізує дисперсію всередині кластерів на кожному кроці об\'єднання.')
    add_body(doc, 'Порівняння з k-means: K-means значно швидший (O(n)), тоді як ієрархічна кластеризація вимагає побудови матриці відстаней, що займає O(n³). Ієрархічна кластеризація часто точніша на складних структурах даних, оскільки вона не припускає сферичної форми кластерів на ранніх етапах. K-means дуже чутливий до викидів (оскільки центроїд зміщується), тоді як ієрархічна кластеризація просто виділить шумові точки в окремі маленькі гілки дерева.')
    add_image_with_caption(doc, os.path.join(lr_dir, 'Lab5_Dendrogram.png'), 'Рисунок 2 – Дендрограма ієрархічної кластеризації')

    # --- 3. t-SNE ---
    add_heading_centered(doc, '3. Зниження розмірності (t-SNE) у 3D просторі', level=2, font_size=14)
    add_body(doc, 't-SNE було використано для нелінійного зниження розмірності з 4D (всі оригінальні ознаки Ірису) до 3D. t-SNE ідеально вловлює локальні структури. На отриманому графіку можна спостерігати три яскраво виражені хмари точок у тривимірному просторі. На відміну від k-means, який просто застосовує математичну відстань, t-SNE використовує ймовірнісні розподіли, що дозволяє "розтягнути" кластери Versicolor та Virginica один від одного набагато якісніше.')
    add_image_with_caption(doc, os.path.join(lr_dir, 'Lab5_tSNE_3D.png'), 'Рисунок 3 – t-SNE 3D візуалізація Iris')

    # --- 4. DBSCAN ---
    add_heading_centered(doc, '4. DBSCAN та PCA', level=2, font_size=14)
    add_body(doc, 'DBSCAN (Density-Based Spatial Clustering of Applications with Noise) будує кластери на основі щільності точок. Щоб цей алгоритм адекватно працював на датасеті Ірисів, спочатку було застосовано PCA (Метод головних компонент) для зниження розмірності до 2.')
    add_body(doc, 'Аналіз параметрів: eps (радіус) = 0.6, при стандартних 0.5 два класи (Versicolor/Virginica) зливаються в один суцільний кластер. min_samples = 4, оскільки датасет складається всього зі 150 точок.')
    add_body(doc, 'DBSCAN чудово підходить для даних довільної форми (не обов\'язково сферичних, як у k-means). Його головна перевага — він автоматично знаходить шумові точки (викиди) і не присвоює їх жодному кластеру. Однак для датасету Ірисів з перехрещеною щільністю його налаштування набагато складніше, ніж просте використання k-means або agglomerative clustering.')
    add_image_with_caption(doc, os.path.join(lr_dir, 'Lab5_DBSCAN_PCA.png'), 'Рисунок 4 – DBSCAN з PCA (eps=0.6, min_samples=4)')

    # --- Висновки ---
    add_heading_centered(doc, '5. Висновки', level=2, font_size=14)
    add_body(doc, 'У ході лабораторної роботи було досліджено чотири різних методи навчання без вчителя на датасеті Ірисів Фішера (150 зразків, 4 ознаки, 3 класи).')
    add_body(doc, '1. K-means показав найкращі результати на парі ознак пелюстки (Petal length / Petal width), де кластери мають майже сферичну форму. Алгоритм працює швидко, але потребує заздалегідь вказати кількість кластерів k та чутливий до викидів.')
    add_body(doc, '2. Ієрархічна кластеризація (метод Ward) побудувала дендрограму, яка наочно демонструє послідовне об\'єднання точок у кластери. Цей метод не потребує знання k наперед, але має обчислювальну складність O(n\u00b3), що робить його непридатним для великих датасетів.')
    add_body(doc, '3. t-SNE дозволив візуалізувати 4-вимірні дані у 3D просторі зі збереженням локальних відстаней. Три класи ірисів чітко розділяються у тривимірному просторі, що підтверджує наявність трьох різних підмножин у даних.')
    add_body(doc, '4. DBSCAN з PCA ефективно виявив кластери довільної форми та автоматично позначив аномальні точки як шум (Outliers). Підбір гіперпараметрів eps та min_samples є критичним для якості кластеризації.')
    add_body(doc, 'Загалом, кожен метод має свої сильні та слабкі сторони, і вибір конкретного алгоритму залежить від характеру даних, їх обсягу та вимог до інтерпретованості результатів.')

    insert_title_page(doc, 5)
    out = os.path.join(lr_dir, 'Lab5_Report.docx')
    doc.save(out)
    print(f'[OK] {out}')


# ─────────────────────────────────────────────
if __name__ == '__main__':
    print('Генерація .docx звітів...\n')
    generate_lab1()
    generate_lab2()
    generate_lab3()
    generate_lab4()
    generate_lab5()
    print('\n[OK] Усі 5 звітів успішно створено!')
